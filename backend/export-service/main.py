from fastapi import FastAPI, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel

app = FastAPI(title="Export Service")

class ExportRequest(BaseModel):
    report: str
    format: str = "docx"

@app.post("/api/export")
async def export_report(request: ExportRequest):
    if request.format == "docx":
        try:
            from docx import Document
            from io import BytesIO

            doc = Document()
            for line in request.report.split("\n"):
                stripped = line.strip()
                if stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped:
                    doc.add_paragraph(stripped)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return Response(
                content=buffer.read(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=research_report.docx"}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx not installed")

    elif request.format == "pdf":
        try:
            from fpdf import FPDF
            from io import BytesIO

            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Helvetica", size=11)

            for line in request.report.split("\n"):
                stripped = line.strip()
                if stripped.startswith("## "):
                    pdf.set_font("Helvetica", "B", 14)
                    pdf.cell(0, 10, stripped[3:], new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", size=11)
                elif stripped.startswith("# "):
                    pdf.set_font("Helvetica", "B", 16)
                    pdf.cell(0, 12, stripped[2:], new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", size=11)
                elif stripped:
                    pdf.multi_cell(0, 6, stripped)
                else:
                    pdf.ln(4)

            buffer = BytesIO()
            pdf.output(buffer)
            buffer.seek(0)
            return Response(
                content=buffer.read(),
                media_type="application/pdf",
                headers={"Content-Disposition": "attachment; filename=research_report.pdf"}
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="fpdf2 not installed")
    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Use 'docx' or 'pdf'.")
