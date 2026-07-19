"""
ARS Orchestrator Service — Main Entrypoint

FastAPI application that replaces both agent-service and search-service
with a graph-orchestrated research system.
"""

from __future__ import annotations

import os
import json
import logging
from contextlib import asynccontextmanager
from typing import AsyncGenerator

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel, Field
from dotenv import load_dotenv

# Load environment
load_dotenv()
env_path = ".env.example"
if os.path.exists(env_path):
    load_dotenv(dotenv_path=env_path)

# ── Observability (must initialize before other imports) ──────────
from observability.logging import setup_logging
from observability.tracing import setup_tracing

setup_logging(
    level=os.getenv("LOG_LEVEL", "INFO"),
    format=os.getenv("LOG_FORMAT", "console"),
)

log = logging.getLogger("ars.main")

# ── Internal imports ──────────────────────────────────────────────
from db.connection import get_pool, close_pool, init_schema
from db.migrations import run_migrations
from state.state_manager import StateManager
from graph.engine import GraphEngine
from tools.registry import ToolRegistry
from tools.llm import LLMTool
from tools.retrieval import RetrievalTool, KnowledgeStatsTool
from tools.citation import CitationVerificationTool
from tools.memory import StoreMemoryTool, RecallMemoryTool
from tools.search import SearchWebTool
from tools.ingestion import IngestionTool
from knowledge.crawler import KnowledgeCrawler
from observability.metrics import MetricsCollector


# ── Lifespan ──────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application startup and shutdown."""
    log.info("Starting ARS Orchestrator Service...")

    # Initialize database
    await get_pool()
    await init_schema()
    await run_migrations()
    log.info("Database initialized ✓")

    # Initialize tool registry
    registry = ToolRegistry()
    registry.register(LLMTool())
    registry.register(RetrievalTool())
    registry.register(KnowledgeStatsTool())
    registry.register(CitationVerificationTool())
    registry.register(StoreMemoryTool())
    registry.register(RecallMemoryTool())
    registry.register(SearchWebTool())
    registry.register(IngestionTool())
    log.info("Tool registry initialized: %d tools ✓", len(registry))

    # Store in app state for access in endpoints
    app.state.registry = registry
    app.state.state_manager = StateManager()
    app.state.graph_engine = GraphEngine(
        state_manager=app.state.state_manager,
        tool_registry=registry,
        max_parallel=int(os.getenv("MAX_PARALLEL_WORKERS", "3")),
    )
    app.state.crawler = KnowledgeCrawler(llm_tool=registry.get("call_llm"))
    app.state.metrics = MetricsCollector()

    log.info("ARS Orchestrator Service ready ✓")
    yield

    # Shutdown
    await close_pool()
    log.info("ARS Orchestrator Service shutdown complete")


app = FastAPI(
    title="ARS Orchestrator",
    description="Graph-orchestrated autonomous research system",
    version="2.0.0",
    lifespan=lifespan,
)
setup_tracing(app=app)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Request/Response Models ───────────────────────────────────────

class CreateTaskRequest(BaseModel):
    message: str
    provider: str = "gemini"
    agent_providers: dict[str, str] = Field(default_factory=dict)
    agent_temperatures: dict[str, float] = Field(default_factory=dict)
    library_only: bool = False
    library_papers: list[dict] = Field(default_factory=list)


class CrawlRequest(BaseModel):
    queries: list[str]


class ExportRequest(BaseModel):
    report: str
    format: str = "docx"


# ── Health & Settings ─────────────────────────────────────────────

@app.get("/health")
async def health_check():
    return {"status": "ok", "service": "orchestrator", "version": "2.0.0"}


@app.get("/api/settings")
async def get_settings():
    """Return available providers."""
    llm_tool: LLMTool = app.state.registry.get("call_llm")

    # Check provider availability
    providers = {}
    for provider_name in ["gemini", "openai", "local_llama"]:
        try:
            if provider_name == "gemini":
                from tools.llm import _get_gemini_client
                _get_gemini_client()
                providers[provider_name] = {"available": True}
            elif provider_name == "openai":
                from tools.llm import _get_openai_client
                _get_openai_client()
                providers[provider_name] = {"available": True}
            elif provider_name == "local_llama":
                from tools.llm import _get_local_llama_client
                _get_local_llama_client()
                providers[provider_name] = {"available": True}
        except Exception:
            providers[provider_name] = {"available": False}

    return {"providers": providers}


@app.get("/api/health/local_llama")
async def check_local_llama_health():
    """Probe the local LLaMA server."""
    import httpx

    base_url = os.getenv("LOCAL_LLAMA_URL", "http://host.docker.internal:8080/v1")
    try:
        health_url = base_url.rstrip("/").rsplit("/v1", 1)[0] + "/health"
        async with httpx.AsyncClient(timeout=1.5) as client:
            probe = await client.get(health_url)
            return {
                "available": probe.status_code == 200,
                "base_url": base_url,
                "model": os.getenv("LOCAL_LLAMA_MODEL", "local-model"),
            }
    except Exception as e:
        return {"available": False, "reason": str(e)}


# ── Task Management ──────────────────────────────────────────────

@app.post("/api/chat")
async def create_and_run_task(request: CreateTaskRequest):
    """
    Create a new research task and stream its execution via SSE.

    This is the primary endpoint — backward-compatible with the
    existing frontend's /api/chat contract.
    """
    if not request.message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    sm: StateManager = app.state.state_manager
    engine: GraphEngine = app.state.graph_engine

    # Create task
    state = await sm.create_task(
        query=request.message.strip(),
        provider=request.provider,
        agent_providers=request.agent_providers,
        agent_temperatures=request.agent_temperatures,
        library_only=request.library_only,
        library_papers=request.library_papers,
    )

    log.info("Task created: %s — %r", state.task_id, request.message[:80])

    # Stream execution
    return StreamingResponse(
        engine.execute(state.task_id),
        media_type="text/event-stream",
    )


@app.post("/api/tasks")
async def create_task(request: CreateTaskRequest):
    """Create a task without starting it (for async execution)."""
    if not request.message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    sm: StateManager = app.state.state_manager
    state = await sm.create_task(
        query=request.message.strip(),
        provider=request.provider,
        agent_providers=request.agent_providers,
    )
    return {"task_id": state.task_id, "status": state.status.value}


@app.get("/api/tasks/{task_id}")
async def get_task(task_id: str):
    """Get current task state."""
    sm: StateManager = app.state.state_manager
    try:
        state = await sm.load_task(task_id)
        return state.model_dump()
    except ValueError:
        raise HTTPException(status_code=404, detail="Task not found")


@app.get("/api/tasks/{task_id}/stream")
async def stream_task(task_id: str):
    """Stream a task's execution (for tasks created via POST /api/tasks)."""
    engine: GraphEngine = app.state.graph_engine
    return StreamingResponse(
        engine.execute(task_id),
        media_type="text/event-stream",
    )


@app.post("/api/tasks/{task_id}/resume")
async def resume_task(task_id: str):
    """Resume a failed/paused task from the latest checkpoint."""
    sm: StateManager = app.state.state_manager
    engine: GraphEngine = app.state.graph_engine

    try:
        state = await sm.restore_from_checkpoint(task_id)
        await sm.save_task(state)
    except ValueError:
        raise HTTPException(status_code=404, detail="No checkpoint found")

    return StreamingResponse(
        engine.execute(task_id),
        media_type="text/event-stream",
    )


@app.get("/api/tasks")
async def list_tasks(status: str | None = None, limit: int = 50):
    """List tasks with optional status filter."""
    sm: StateManager = app.state.state_manager
    from models.task_state import TaskStatus

    status_filter = None
    if status:
        try:
            status_filter = TaskStatus(status)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid status: {status}")

    tasks = await sm.list_tasks(status=status_filter, limit=limit)
    return {"tasks": tasks}


# ── Evaluation & Metrics ─────────────────────────────────────────

@app.get("/api/tasks/{task_id}/evaluation")
async def get_evaluation(task_id: str):
    """Get evaluation metrics for a completed task."""
    sm: StateManager = app.state.state_manager
    try:
        state = await sm.load_task(task_id)
        if state.metrics:
            return state.metrics.model_dump()
        return {"error": "No metrics available (task may not be complete)"}
    except ValueError:
        raise HTTPException(status_code=404, detail="Task not found")


@app.get("/api/metrics")
async def system_metrics():
    """Get system-wide metrics."""
    collector: MetricsCollector = app.state.metrics
    return await collector.get_system_metrics()


# ── Knowledge Acquisition ────────────────────────────────────────

@app.post("/api/knowledge/crawl")
async def trigger_crawl(request: CrawlRequest):
    """Trigger a knowledge crawl (independent of tasks)."""
    crawler: KnowledgeCrawler = app.state.crawler

    async def stream_crawl():
        result = await crawler.crawl(request.queries)
        yield f"data: {json.dumps({'type': 'crawl_complete', 'result': {'papers': len(result.papers), 'ingested': len(result.ingestion_results), 'duration': result.duration_seconds}})}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(stream_crawl(), media_type="text/event-stream")


@app.get("/api/knowledge/stats")
async def knowledge_stats():
    """Get knowledge base statistics."""
    stats_tool = app.state.registry.get("knowledge_stats")
    result = await stats_tool()
    return result.data if result.success else {"error": result.error}


# ── Paper Analysis (preserved from search-service) ───────────────

@app.get("/api/papers/search")
async def search_papers(q: str):
    """Search for papers via the knowledge crawler."""
    if not q.strip():
        raise HTTPException(status_code=400, detail="Query required")

    crawler: KnowledgeCrawler = app.state.crawler

    async def stream_search():
        import asyncio
        queue = asyncio.Queue()

        async def sse_callback(event_json: str):
            await queue.put(f"data: {event_json}\n\n")

        async def run_crawler():
            try:
                await queue.put(f"data: {json.dumps({'type': 'agent_step', 'step': 0, 'message': 'Starting headless browser...', 'status': 'running'})}\n\n")
                result = await crawler.crawl([q.strip()], sse_callback=sse_callback)
                for paper in result.papers:
                    await queue.put(f"data: {json.dumps({'type': 'paper_found', 'paper': paper})}\n\n")
                await queue.put(f"data: {json.dumps({'type': 'final_result', 'papers': result.papers, 'count': len(result.papers)})}\n\n")
            except Exception as e:
                await queue.put(f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n")
            finally:
                await queue.put("data: [DONE]\n\n")
                await queue.put(None)

        task = asyncio.create_task(run_crawler())

        try:
            while True:
                item = await queue.get()
                if item is None:
                    break
                yield item
        finally:
            if not task.done():
                task.cancel()

    return StreamingResponse(stream_search(), media_type="text/event-stream")


@app.post("/api/papers/analyze")
async def analyze_papers(request: dict):
    """Analyze uploaded papers or a single URL."""
    papers = request.get("papers", [])
    if "url" in request:
        papers.append({"title": "Imported by URL", "pdfUrl": request["url"], "id": request["url"]})

    if not papers:
        raise HTTPException(status_code=400, detail="No papers or URL provided")

    ingestion_tool = app.state.registry.get("ingest_document")
    results = []
    for paper in papers:
        result = await ingestion_tool(
            title=paper.get("title", "Unknown"),
            pdf_url=paper.get("pdfUrl", ""),
            paper_id=paper.get("id", ""),
            arxiv_id=paper.get("id", ""),
        )
        results.append(result.data if result.success else {"error": result.error})

    return {"results": results}


# ── Export (preserved from original API) ──────────────────────────

@app.post("/api/export")
async def export_report(request: ExportRequest):
    """Export the report to DOCX or PDF format."""
    if request.format == "docx":
        try:
            from docx import Document
            from io import BytesIO

            doc = Document()
            for line in request.report.split("\n"):
                stripped = line.strip()
                if stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped:
                    doc.add_paragraph(stripped)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return Response(
                content=buffer.read(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=research_report.docx"},
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx not installed")

    elif request.format == "pdf":
        try:
            from fpdf import FPDF
            from io import BytesIO

            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Helvetica", size=11)

            for line in request.report.split("\n"):
                stripped = line.strip()
                if stripped.startswith("## "):
                    pdf.set_font("Helvetica", "B", 14)
                    pdf.cell(0, 10, stripped[3:], new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", size=11)
                elif stripped.startswith("# "):
                    pdf.set_font("Helvetica", "B", 16)
                    pdf.cell(0, 12, stripped[2:], new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", size=11)
                elif stripped:
                    pdf.multi_cell(0, 6, stripped)
                else:
                    pdf.ln(4)

            buffer = BytesIO()
            pdf.output(buffer)
            buffer.seek(0)
            return Response(
                content=buffer.read(),
                media_type="application/pdf",
                headers={"Content-Disposition": "attachment; filename=research_report.pdf"},
            )
        except ImportError:
            raise HTTPException(status_code=500, detail="fpdf2 not installed")

    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Use 'docx' or 'pdf'.")
